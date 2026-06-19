"""
Workflow B — English courses supplement.
Reuses the login + navigation from Workflow A, then extracts
English section text, translates it to Spanish, builds a
formatted Word document, and prints it.
"""

import os
import subprocess
import platform
import deepl
from playwright.sync_api import sync_playwright, Page
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils import get_credentials, get_deepl_key, find_weekly_report_url, login

OUTPUT_DIR = "output"
DOCX_FILENAME = "english_supplement_es.docx"

# English course labels to look for in the report
ENGLISH_SUBJECTS = ["GRAMMAR", "READING", "SPELLING", "ORAL", "SCIENCE"]


# Day sort order for Spanish weekday names
DAY_ORDER = ["lunes", "martes", "miércoles", "jueves", "viernes"]


def extract_english_sections(page: Page) -> dict[str, list[dict]]:
    sections = {}
    current_day = "Sin día"

    rows = page.locator("tr.Row, tr.GroupCaption").all()
    for row in rows:
        if "GroupCaption" in (row.get_attribute("class") or ""):
            current_day = row.inner_text().strip()
            continue

        cells = row.locator("td").all()
        if len(cells) < 3:
            continue

        subject = cells[0].inner_text().strip().upper()
        if subject in ENGLISH_SUBJECTS:
            # ← capture Asignación Tipo
            tipo = cells[1].inner_text().strip()
            description = cells[2].inner_text().strip()
            if subject not in sections:
                sections[subject] = []
            sections[subject].append({
                "day": current_day,
                "tipo": tipo,
                "description": description
            })

    if sections:
        found = ", ".join(sections.keys())
        total = sum(len(v) for v in sections.values())
        print(f"  ✓ Extracted {total} entries across: {found}")
    else:
        print("  ✗ No English sections found — check selectors")

    return sections


def translate_sections(sections: dict[str, list[dict]], api_key: str) -> dict[str, list[dict]]:
    translator = deepl.Translator(api_key)
    translated = {}

    for subject, entries in sections.items():
        translated[subject] = []
        for entry in entries:
            result = translator.translate_text(
                entry["description"], target_lang="ES")
            translated[subject].append({
                "day": entry["day"],
                "tipo": entry["tipo"],           # ← pass through unchanged
                "description": result.text
            })
        print(f"  ✓ Translated: {subject} ({len(entries)} entries)")

    return translated


# Assignment types that get red + bold treatment
EMPHASIS_TIPOS = {"exam", "examen", "homework", "tarea", "practice",
                  "práctica", "practicas", "prácticas", "test", "quiz"}


def build_document(sections: dict[str, list[dict]], output_path: str) -> None:
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading("Tareas de Inglés — Semana Actual", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
    doc.add_paragraph()

    # ── Reorganize by day ───────────────────────────────────────────────────
    by_day = {}
    for subject, entries in sections.items():
        for entry in entries:
            day = entry["day"]
            if day not in by_day:
                by_day[day] = []
            by_day[day].append({
                "subject": subject,
                "tipo": entry["tipo"],
                "description": entry["description"]
            })

    # ── Sort days by weekday order ──────────────────────────────────────────
    def day_sort_key(day_str):
        day_lower = day_str.lower()
        for i, name in enumerate(DAY_ORDER):
            if name in day_lower:
                return i
        return 99

    sorted_days = sorted(by_day.keys(), key=day_sort_key)

    # ── Render each day ─────────────────────────────────────────────────────
    for i, day in enumerate(sorted_days):
        entries = by_day[day]

        # Divider line between days (not before the first one)
        if i > 0:
            divider = doc.add_paragraph()
            divider.runs  # ensure paragraph exists
            divider.paragraph_format.space_before = Pt(4)
            divider.paragraph_format.space_after = Pt(4)
            pPr = divider._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '6')
            top.set(qn('w:space'), '1')
            top.set(qn('w:color'), 'CCCCCC')
            pBdr.append(top)
            pPr.append(pBdr)

        # Day heading
        day_heading = doc.add_heading(day.capitalize(), level=2)
        day_heading.runs[0].font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)

        for entry in entries:
            is_emphasis = entry["tipo"].lower() in EMPHASIS_TIPOS

            # Subject + tipo on same line
            label_para = doc.add_paragraph()
            subject_run = label_para.add_run(entry["subject"].capitalize())
            subject_run.bold = True
            subject_run.font.size = Pt(11)
            subject_run.font.color.rgb = (
                RGBColor(0xFF, 0x00, 0x00) if is_emphasis
                else RGBColor(0x53, 0x4A, 0xB7)
            )

            tipo_run = label_para.add_run(f"  [{entry['tipo']}]")
            tipo_run.bold = is_emphasis
            tipo_run.italic = True
            tipo_run.font.size = Pt(10)
            tipo_run.font.color.rgb = (
                RGBColor(0xFF, 0x00, 0x00) if is_emphasis
                else RGBColor(0x88, 0x88, 0x88)
            )

            # Description
            para = doc.add_paragraph(entry["description"])
            para.runs[0].font.size = Pt(11)
            if is_emphasis:
                para.runs[0].bold = True

    doc.save(output_path)
    print(f"  ✓ Document saved → {output_path}")


def print_docx(docx_path: str) -> None:
    abs_path = os.path.abspath(docx_path)
    print(f"  → Opening in Edge for review...")

    edge_paths = [
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    ]

    edge = next((p for p in edge_paths if os.path.exists(p)), None)
    if not edge:
        raise FileNotFoundError("Edge not found")

    subprocess.Popen([edge, abs_path])
    print("  ✓ Opened in Edge")


def run_workflow_b() -> None:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, DOCX_FILENAME)
    user, password = get_credentials()
    deepl_key = get_deepl_key()

    with sync_playwright() as p:
        # set True once confirmed working
        browser = p.chromium.launch(headless=False)
        page = browser.new_page()

        try:
            login(page, user, password)
            report_url = find_weekly_report_url(page)

            if not report_url:
                print("  ✗ Workflow B aborted: no report URL found")
                return

            page.goto(report_url)
            page.wait_for_load_state("networkidle")

            sections = extract_english_sections(page)
            if not sections:
                print("  ✗ Workflow B aborted: no English sections extracted")
                return

            translated = translate_sections(sections, deepl_key)
            build_document(translated, output_path)
            print_docx(output_path)

        finally:
            browser.close()
